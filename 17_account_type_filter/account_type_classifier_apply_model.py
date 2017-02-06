# -*- coding: utf-8 -*-
from __future__ import print_function, division

#------------------------------------------------------------------------------
# TWITTER ACCOUNT TYPE CLASSIFIER - APPLY MODEL
#------------------------------------------------------------------------------

# Functions:
# [1] ...

# Version: 2.0
# Last edited: 21 Dec 2016
# Edited by: Minh PHAN

#------------------------------------------------------------------------------
# Global variables and settings
#------------------------------------------------------------------------------

# Seting working directory
import os
os.chdir('D:\\SentimentTM\\17_account_type_filter')

# Variables
n_dimension = 300 # Word2Vec dimensions

#------------------------------------------------------------------------------
# Initiating
#------------------------------------------------------------------------------

# Essential packages
import pandas as pd
from time import time

# Other functional packages
import pickle

#------------------------------------------------------------------------------
# Functions to read pre-processed files
#------------------------------------------------------------------------------

# Function to create a list of image features names
def findFeaturesFeatureNames():
    f_names = ['contrast','avg_h','avg_s', 'avg_v', 'pleasure', 'arousal', 'dominance']

    f_names = f_names + ['hue_hist_' + str(i) for i in range(12)]
    f_names = f_names + ['saturation_hist_' + str(i) for i in range(5)]
    f_names = f_names + ['brightness_hist_' + str(i) for i in range(3)]
    f_names = f_names + ['std_hueHist', 'std_satHist', 'std_brHist']
    f_names = f_names + ['haralicks_f_' + str(i) for i in range(13)]
    return f_names
    
# Function to create a list of text vector features names
def word2vec_features_names():
    f_names = []
    for i in range(n_dimension):
        f_names = f_names + ['w2v_' + str(i)]
    return f_names

#------------------------------------------------------------------------------
# MAIN: Apply the pre-train xgboost model on new data
#------------------------------------------------------------------------------

# Load model
file_in = '.\\model\\output\\17_twitter_account_type_classifier_model.pkl'
le, clf_model = pickle.load(open(file_in, "rb"))

# Load pre-processed data
file_in = '.\\input\\16_tweets_location_full_features.tsv'
userDataWithF = pd.read_csv(file_in, sep='\t', encoding='utf-8')

# Apply models
t0 = time()
featureCols = findFeaturesFeatureNames() + word2vec_features_names()
X = userDataWithF[featureCols]

y_pred_proba = clf_model.predict_proba(X) # 1 = org / 0 = idv
y_pred = clf_model.predict(X) # 1 = org / 0 = idv
print('Running time:', time()-t0)

# Save data
userDataWithF['account_type_pred'] = le.inverse_transform(y_pred)
userDataWithF['account_type_pred_confidence'] = y_pred_proba.max(axis=1)

file_out = '.\\output\\17_twitter_account_type_classified.tsv'
userDataWithF.to_csv(file_out, encoding='utf-8', sep='\t', index=False)

# Save only individual tweets
#userDataWithF_idv = userDataWithF[userDataWithF['account_type_pred'] == 0]
#file_out = '.\\output\\17_twitter_account_type_individual_tweets.tsv'
#userDataWithF_idv.to_csv(file_out, encoding='utf-8', sep='\t', index=False)

#------------------------------------------------------------------------------
# Print out the output [individual] idv = 0; org = 1;
#------------------------------------------------------------------------------

from docx import Document
from docx.shared import Inches

imgs_dir = '..\\16_enrich_tweets_data\\output\\16_twitter_profile_images'
userDataWithF['pred_proba'] = y_pred_proba[:,1]

# Extract all accounts that being predicted as individual
ind_df = userDataWithF.loc[userDataWithF['account_type_pred']=='idv']

threshold = 0.50 # This is the threshold for decision function of the prediction
ind_df = ind_df.loc[ind_df['pred_proba'] <= threshold] # idv = 0
ind_df.sort('account_type_pred_confidence', ascending=False, inplace=True)
ind_df.reset_index(drop=True, inplace=True)
ind_df.fillna('< Missing >', inplace=True)

# Create a Word documents
def df2docx(df_input, header, file_out):
    document = Document()
    document.add_heading(header, 1)
    
    # Create a table to hold the result
    table = document.add_table(1, cols=3)
    table.style = 'TableGrid'
    
    # Set columns width
    col_width = [1.5, 5.0, 0.5]
    table.rows[0].cells[0].width = Inches(col_width[0])
    table.rows[0].cells[1].width = Inches(col_width[1])
    table.rows[0].cells[2].width = Inches(col_width[2])
    
    # Set columns titles
    table.rows[0].cells[0].paragraphs[0].add_run('Avatar').bold = True
    table.rows[0].cells[1].paragraphs[0].add_run('Information').bold = True
    table.rows[0].cells[2].paragraphs[0].add_run('Confidence').bold = True
    
    # For loop to add results
    count = 0
    for index, row in df_input.iterrows():
    
        # First thing to do, add new row    
        table.add_row()
        count += 1
        #print(count) # To track process
        table.rows[count].cells[0].width = Inches(col_width[0])
        table.rows[count].cells[1].width = Inches(col_width[1])
        table.rows[count].cells[2].width = Inches(col_width[2])
    
        # Add the picture
        run = table.rows[count].cells[0].paragraphs[0].add_run()
        file_name = os.path.join(imgs_dir, str(row['user_id_str']) + '.jpg')
        try:
            run.add_picture(file_name, width=Inches(1.25))
        except:
            table.rows[count].cells[0].add_paragraph('< Missing >')
        
        # Add profile information
        table.rows[count].cells[1].add_paragraph('Name: ' + row['user_screen_name'] + ' / ' + row['user_name'] + ' / ' + str(row['user_id_str']) + '\n')
        table.rows[count].cells[1].add_paragraph('Location: ' + row['user_location'] + '\n')
        table.rows[count].cells[1].add_paragraph('Description: ' + row['user_description'] + '\n')
        table.rows[count].cells[1].add_paragraph('Tweet: ' + row['text'] + '\n')
        
        # Add predict probability
        table.rows[count].cells[2].add_paragraph('{0:0.2f}'.format(row['account_type_pred_confidence']))
        
        #if count == 100: break # To debug only
    
    # Save the documents to file
    document.save(file_out)

file_out = '.\\output\\individual_accounts_top100.docx'
df2docx(ind_df.head(100), header='INDIVIDUAL TWEETS - TOP 100', file_out=file_out)

file_out = '.\\output\\individual_accounts_bottom100.docx'
df2docx(ind_df.tail(100), header='INDIVIDUAL TWEETS - BOTTOM 100', file_out=file_out)

#------------------------------------------------------------------------------
# Print out the output [organizational]
#------------------------------------------------------------------------------

# Extract all accounts that being predicted as organization
org_df = userDataWithF.loc[userDataWithF['account_type_pred']=='org']

threshold = 0.50 # This is the threshold for decision function of the prediction
org_df = org_df.loc[org_df['pred_proba'] > threshold] # org = 1
org_df.sort('account_type_pred_confidence', ascending=False, inplace=True)
org_df.reset_index(drop=True, inplace=True)
org_df.fillna('< Missing >', inplace=True)

file_out = '.\\output\\organization_accounts_top100.docx'
df2docx(org_df.head(100), header='ORGANIZATIONAL TWEETS - TOP 100', file_out=file_out)

file_out = '.\\output\\organization_accounts_bottom100.docx'
df2docx(org_df.tail(100), header='ORGANIZATIONAL TWEETS - BOTTOM 100', file_out=file_out)

#------------------------------------------------------------------------------